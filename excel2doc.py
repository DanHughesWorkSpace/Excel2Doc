from docxtpl import DocxTemplate;
import os
from xlrd import open_workbook
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Inches, Pt

# IF UPDATING THE EXCEL FORMAT :
    #  CONFIGURE THE section_dict  DICTIONARY AND LEAVE REMAINING CODE AS IS
    #  YOU WILL HAVE TO UPDATE docvalues
wb = open_workbook('format.xls')
for s in wb.sheets():
    docvalues = []
    for row in range(s.nrows):
        col_value = []
        for col in range(s.ncols):
            value  = (s.cell(row,col).value)
            
            try : value = str((value))
            except : pass
            if len(value) > 1 and value[0] != 'Q' and value[2] == '0':
                value = value[0]
                col_value.append(value)
            else:
                col_value.append(value)
        docvalues.append(col_value)

context = {}

applicationName = docvalues[0][3]
context["applicationName"] = applicationName
applicationPrefix = docvalues[0][4]
context["applicationPrefix"] = applicationPrefix
documentId = docvalues[0][5]
context["documentId"] = documentId
documentTitle = docvalues[0][6]
context["documentTitle"] = documentTitle
documentRevision = docvalues[0][7]
context["documentRevision"] = documentRevision

section_dict = {}

section_array = []  
section = 1
i = 0;
for item in docvalues:
    if item[0] == str(section):
        section_array.append(item)    
    else:
        section_dict[section] = section_array
        section_array = []
        section_array.append(item)
        section = section + 1
        section_dict[section] = section_array

def creater_section_header(section):
        section_header = section[0][2]
        tpl.add_heading(section_header, level=1)

def create_table(section):
        if len(section) > 0:
            table = tpl.add_table(1,4)
            table.rows[0].cells[0].text = "Function ID"
            table.rows[0].cells[1].text = "Description"
            table.rows[0].cells[2].text = "Priority"
            table.rows[0].cells[3].text = "Criticality"
            table.rows[0].cells[0].width = Inches(1)
            table.rows[0].cells[1].width = Inches(5)
            table.rows[0].cells[2].width = Inches(1)
            table.rows[0].cells[3].width = Inches(1)
            row = table.rows[0]
            for cell in row.cells:
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="0C3C60"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_2)
            i = 0
            while i < len(section):
                new_row = table.add_row().cells
                new_row[0].text = section[i][2]
                new_row[1].text = section[i][3]
                new_row[2].text = section[i][4]
                new_row[3].text = section[i][5]
                i = i + 1

def populate_section(section):
        i = 0
        table_content = []
        while i < len(section):
            for char in section[i][1]:
                if char == ".":
                    tpl.add_heading(section[i][2], level=2)
            if section[i][1] == "text":
                tpl.add_paragraph(section[i][2])
            elif section[i][1] == "table":
                table_content.append(section[i])
            i = i + 1
        create_table(table_content)

tpl = DocxTemplate("Word_template/excel2doc_template.docx")

for key in section_dict:
    section = section_dict[key]
    creater_section_header(section)
    populate_section(section)

tpl.render(context)
tpl.save('new_test.docx')

os.startfile("C:/Users/Dan/Documents/projects/excel_2_doc/excel-2-doc/new_test.docx")